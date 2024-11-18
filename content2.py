#content.py
import streamlit as st
from test2 import test_content, get_score
import sqlite3
import contextlib
import io
import docx

def get_content(module_number):
    doc = docx.Document(f"Modules\Module-{module_number}.docx")
    content = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            content += f"<h{level}><b>{para.text}</b></h{level}>"
        else:
            content += f"<p>{para.text}</p>"
    return content

def evaluate_practice_answer(question, answer, output,topics):
    from llm3 import llama3
    review = llama3(
    """
    You are an experienced Python programming tutor and a coding expert. 
    I will provide you with the following details:
    1. A Python coding question.
    2. My current answer to the question.
    3. The output I am getting from running the code.
    4. A list of Python topics I have already completed.

    Your task is to:
    - Carefully review my answer and determine whether it is correct based on the coding question.
    - If my answer is incorrect, offer constructive hints to guide me towards the right solution.
    - If my answer is correct, suggest alternative approaches to solve the problem that could improve efficiency or clarity, 
      but avoid mentioning topics or methods that I have not yet completed and never suggest the user to use f-strings.
    - Motivate me to continue improving my coding skills by giving positive, encouraging feedback.

    Important:
    - Focus on topics I have already covered, and do not suggest anything outside my current knowledge base.
    """,
    f"The coding question is: {question}\nMy answer: \n {answer}\nThe output I'm getting: \n {output}\nThe topics I have completed so far: {topics}\nPlease provide your expert review and guidance accordingly.")
    return review


def module_content(module_number, practice_question,topics_completed):
    st.title(f"Module-{module_number}")

    placeholder = "Your answer here"
    content = get_content(module_number)
    st.markdown(content, unsafe_allow_html=True)

    st.markdown("<h2>Practice Question</h2>", unsafe_allow_html=True)
    st.markdown(f"<h3>{practice_question}</h3>", unsafe_allow_html=True)

    user_answer = st.text_area(label="Write your answer in the below text area:", placeholder=placeholder)
    user_output = st.text_area(label="Paste your output in the below text area", placeholder="Your output here")

    review = evaluate_practice_answer(practice_question, user_answer, user_output,topics_completed)

    if st.button("Submit answer"):
        st.write(review)





# Function to handle module content
def display_module(module_number, practice_question,topics_completed):
    with sqlite3.connect('testing.db', check_same_thread=False) as conn:
        c = conn.cursor()
        c.execute("SELECT module_progress FROM testing WHERE id = ?", (st.session_state['user_id'],))
        module_progress = c.fetchone()
        if module_progress and module_progress[0] >= module_number:
            module_content(module_number, practice_question, topics_completed)
        else:
            st.error(f"You need to pass the test for Module {module_number - 1} to access this module.")

# Function to handle module tests
def display_test(module_number, test_question):
    if st.session_state['module_progress'] >= module_number:
        test_answer = test_content(module_number, test_question)
        if test_answer is not None:
            score = get_score(test_answer)
            if score ==1:
                st.success("You passed the test!")
                if st.session_state['module_progress'] < module_number + 1:
                    with sqlite3.connect('testing.db', check_same_thread=False) as conn:
                        c = conn.cursor()
                        c.execute("UPDATE testing SET module_progress = ? WHERE id = ?", (module_number + 1, st.session_state['user_id']))
                        conn.commit()
                    st.session_state['module_progress'] = module_number + 1
            else:
                st.error("You need to pass the test to unlock the next module.")
    else:
        st.error(f"You need to pass the test for Module {module_number - 1} to access this test.")



def get_file_content(file_name):
    doc = docx.Document(f"{file_name}.docx")
    content = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            content += f"<h{level}>{para.text}</h{level}>"
        else:
            content += f"<p>{para.text}</p>"
    return content

def project_content(file_name,topics_completed):
    from llm3 import llama3
    description=get_file_content(file_name)
    st.markdown(description,unsafe_allow_html=True)
    role=f"""You are a Python instructor. You have taught students the following topics: {topics_completed}. You have assigned a project: {description}.
      A student will provide their project answer and output. Analyze their work in the context of the project question, motivate the student, 
      and provide hints to help them solve the project using only the topics that have been completed. 
    If the student successfully solves the project, motivate them and suggest ways to improve their solution. """

    code=st.text_area(label="Write your code here",placeholder="Paste your code here:",height=200)

    output=st.text_area(label="Paste your output below",placeholder="Paste your output here:")

    review=llama3(role,f"The project is {description} ,The code is {code} and the output is {output}.Please provide a review")

    if st.button("Submit project"):
        st.write(review)

def get_progress(user_id):
    with sqlite3.connect('testing.db', check_same_thread=False) as conn:
        c = conn.cursor()
        c.execute("SELECT module_progress FROM testing WHERE id=?", (user_id,))
        progress = c.fetchone()
        return progress[0] if progress else None